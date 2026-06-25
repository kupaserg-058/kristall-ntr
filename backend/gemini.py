import asyncio
import io
import json
import logging
import pathlib
import uuid
from difflib import SequenceMatcher

from google import genai
from google.genai import types

from backend import storage
from backend.config import settings
from backend.database import AsyncSessionLocal
from backend.models import Confidence, Direction, DocStatus, Document, VerificationStatus

logger = logging.getLogger("kristall.gemini")

MODEL = "gemini-2.5-flash"
CHUNK_SIZE = 15_000   # символов на чанк
MAX_PER_FEDERAL = 5   # максимум активностей на одно федеральное направление
MAX_RETRIES = 3
RETRY_DELAY = 8

_HIERARCHY_FILE = pathlib.Path(__file__).parent.parent / "ntr_hierarchy_final.json"


# ── Иерархия НТР ─────────────────────────────────────────────────────────────

def _build_hierarchy_text() -> str:
    try:
        data = json.loads(_HIERARCHY_FILE.read_text(encoding="utf-8"))
    except Exception:
        return ""
    lines: list[str] = []
    for d in data.get("directions", []):
        lines.append(f"\n{d['id']}. {d['title']}")
        for c in d.get("children", []):
            lines.append(f"   • {c['title']}")
            for g in c.get("children", []):
                lines.append(f"     – {g['title']}")
    return "\n".join(lines)


_HIERARCHY_TEXT = _build_hierarchy_text()

SYSTEM_INSTRUCTION = f"""Ты — эксперт по анализу региональных стратегий в сфере \
научно-технологического развития России.

Семь федеральных приоритетных направлений НТР (Указ Президента РФ №529 от 18.06.2024):
1. Высокоэффективная и ресурсосберегающая энергетика
2. Превентивная и персонализированная медицина, обеспечение здорового долголетия
3. Высокопродуктивное и устойчивое к изменениям природной среды сельское хозяйство
4. Безопасность получения, хранения, передачи и обработки информации
5. Интеллектуальные транспортные и телекоммуникационные системы, включая автономные транспортные средства
6. Укрепление социокультурной идентичности российского общества и повышение уровня его образования
7. Адаптация к изменениям климата, сохранение и рациональное использование природных ресурсов

Для уточнения что относится к каждому направлению используй иерархию:
{_HIERARCHY_TEXT}

ЗАДАЧА: проанализируй фрагмент документа. Найди конкретные мероприятия, цели и задачи \
региона в сфере науки и технологий. Для каждого определи федеральное направление из 7.

СТРОГИЕ ПРАВИЛА:
- Возвращай только УНИКАЛЬНЫЕ тематические направления — не каждое предложение
- Максимум 5 элементов на весь фрагмент
- Предпочитай конкретные технологические активности общим декларациям
- НЕ дублируй похожие активности
- title должен быть краткой формулировкой (не длиннее 15 слов)

Отвечай СТРОГО в формате JSON без markdown:
{{
  "directions": [
    {{
      "title": "краткое название мероприятия/направления региона",
      "federal_match": "одно из 7 направлений выше или null",
      "fragment": "дословная цитата из документа до 300 символов",
      "confidence": "high | medium | low"
    }}
  ]
}}"""

RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "directions": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "federal_match": {"type": "string", "nullable": True},
                    "fragment": {"type": "string"},
                    "confidence": {"type": "string", "enum": ["high", "medium", "low"]},
                },
                "required": ["title", "fragment", "confidence"],
            },
        }
    },
    "required": ["directions"],
}


# ── Ротация ключей ────────────────────────────────────────────────────────────

class KeyRotator:
    def __init__(self, keys: list[str]) -> None:
        if not keys:
            raise RuntimeError("GEMINI_API_KEY не задан")
        self._clients = [genai.Client(api_key=k) for k in keys]
        self._index = 0
        logger.info("Gemini: %d ключ(ей)", len(self._clients))

    @property
    def current(self) -> genai.Client:
        return self._clients[self._index]

    def rotate(self) -> None:
        self._index = (self._index + 1) % len(self._clients)
        logger.warning("Gemini: переключение на ключ #%d", self._index)


_rotator: KeyRotator | None = None


def get_rotator() -> KeyRotator:
    global _rotator
    if _rotator is None:
        _rotator = KeyRotator(settings.gemini_api_keys)
    return _rotator


# ── Извлечение текста ─────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Извлекает текст из PDF через pymupdf с корректными пробелами."""
    import fitz

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        # "words" mode — восстанавливает пробелы между словами
        words = page.get_text("words", sort=True)
        if not words:
            continue
        # Собираем строки, отслеживая переходы
        lines: list[str] = []
        current_line: list[str] = []
        prev_line_no = -1
        for w in words:
            # w = (x0, y0, x1, y1, word, block_no, line_no, word_no)
            line_no = w[6]
            if line_no != prev_line_no and current_line:
                lines.append(" ".join(current_line))
                current_line = []
            current_line.append(w[4])
            prev_line_no = line_no
        if current_line:
            lines.append(" ".join(current_line))
        pages.append("\n".join(lines))
    doc.close()
    return "\n\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def split_into_chunks(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    paragraphs = [p for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        if current_len + len(para) > chunk_size and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(para)
        current_len += len(para)
    if current:
        chunks.append("\n".join(current))
    return chunks


# ── Анализ чанка ──────────────────────────────────────────────────────────────

async def _analyze_chunk(text: str, chunk_num: int, total: int) -> list[dict]:
    rotator = get_rotator()
    prompt = (
        f"[Фрагмент {chunk_num} из {total}]\n\n"
        f"Текст:\n{text}"
    )
    for attempt in range(MAX_RETRIES):
        try:
            response = await rotator.current.aio.models.generate_content(
                model=MODEL,
                contents=[prompt],
                config=types.GenerateContentConfig(
                    system_instruction=SYSTEM_INSTRUCTION,
                    response_mime_type="application/json",
                    response_schema=RESPONSE_SCHEMA,
                ),
            )
            data = _parse_response(response.text or "")
            dirs = data.get("directions", [])
            logger.info("Чанк %d/%d: %d направлений", chunk_num, total, len(dirs))
            return dirs
        except Exception as exc:
            err = str(exc)
            retryable = any(c in err for c in ("503", "429", "UNAVAILABLE", "RESOURCE_EXHAUSTED"))
            if retryable and attempt < MAX_RETRIES - 1:
                if len(rotator._clients) > 1:
                    rotator.rotate()
                delay = RETRY_DELAY * (attempt + 1)
                logger.warning("Чанк %d retry %d через %ds: %s", chunk_num, attempt + 1, delay, err[:60])
                await asyncio.sleep(delay)
            else:
                logger.error("Чанк %d пропущен: %s", chunk_num, err[:150])
                return []
    return []


# ── Умная дедупликация ────────────────────────────────────────────────────────

def _smart_deduplicate(directions: list[dict]) -> list[dict]:
    """Группирует по federal_match, внутри каждой группы удаляет похожие
    по title (порог сходства 0.6), оставляет не более MAX_PER_FEDERAL."""

    groups: dict[str, list[dict]] = {}
    for d in directions:
        key = d.get("federal_match") or "__unmatched__"
        groups.setdefault(key, []).append(d)

    result: list[dict] = []
    for _federal, items in groups.items():
        unique: list[dict] = []
        for item in items:
            title = (item.get("title") or "").strip().lower()
            is_dup = any(
                SequenceMatcher(None, title, (u.get("title") or "").lower()).ratio() > 0.6
                for u in unique
            )
            if not is_dup:
                unique.append(item)
            if len(unique) >= MAX_PER_FEDERAL:
                break
        result.extend(unique)

    return result


# ── Вспомогательные ───────────────────────────────────────────────────────────

def _parse_response(raw: str) -> dict:
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    text = raw.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = lines[1:] if lines[0].startswith("```") else lines
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return json.loads(text)


def _map_confidence(value: str | None) -> Confidence:
    try:
        return Confidence(value)
    except (ValueError, TypeError):
        return Confidence.medium


# ── Главная функция ───────────────────────────────────────────────────────────

async def analyze_document(
    document_id: uuid.UUID,
    file_bytes: bytes,
    mime_type: str,
    original_name: str,
) -> None:
    async with AsyncSessionLocal() as session:
        document = await session.get(Document, document_id)
        if document is None:
            return

        document.status = DocStatus.analyzing
        await session.commit()

        try:
            # 1. Извлечь текст
            is_pdf = original_name.lower().endswith(".pdf")
            text = extract_text_from_pdf(file_bytes) if is_pdf else extract_text_from_docx(file_bytes)
            storage.delete_temp(document_id)

            if not text.strip():
                raise ValueError("Не удалось извлечь текст из документа")

            logger.info("Документ %s: %d символов, ", document_id, len(text))

            # 2. Чанки
            chunks = split_into_chunks(text, CHUNK_SIZE)
            logger.info("Документ %s: %d чанков", document_id, len(chunks))

            # 3. Анализ
            all_directions: list[dict] = []
            for i, chunk in enumerate(chunks, 1):
                dirs = await _analyze_chunk(chunk, i, len(chunks))
                all_directions.extend(dirs)
                if i < len(chunks):
                    await asyncio.sleep(2)

            # 4. Умная дедупликация
            unique = _smart_deduplicate(all_directions)
            logger.info("Документ %s: %d → %d после дедупликации", document_id, len(all_directions), len(unique))

            # 5. Сохранить
            for item in unique:
                title = (item.get("title") or "").strip()
                if not title:
                    continue
                federal_match = item.get("federal_match")
                if federal_match in (None, "", "null"):
                    federal_match = None
                session.add(Direction(
                    document_id=document.id,
                    title=title,
                    federal_match=federal_match,
                    fragment=(item.get("fragment") or "")[:400],
                    confidence=_map_confidence(item.get("confidence")),
                    verification_status=VerificationStatus.pending,
                ))

            document.status = DocStatus.done
            document.error_message = None
            await session.commit()
            logger.info("Документ %s: %d направлений сохранено", document_id, len(unique))

        except Exception as exc:
            await session.rollback()
            document = await session.get(Document, document_id)
            if document is not None:
                document.status = DocStatus.error
                document.error_message = f"{type(exc).__name__}: {exc}"[:4000]
                await session.commit()
            logger.exception("Ошибка анализа %s", document_id)
